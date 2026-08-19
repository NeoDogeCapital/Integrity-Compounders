"""
trade_briefing.py — advisor-facing trade briefing generator (house template)

The standing template for ALL trade communications to other advisors (Niko,
2026-08-13), modeled on the six-name Model Additions Briefing of the same date.
For each ticker it produces:

  1. "What they do"        — 2-sentence business description   (prose, per-run)
  2. "Catalysts & themes"  — 2-sentence current catalysts       (prose, per-run)
  3. Model findings table  — quad, QGS tier+score, earnings quality, trend +
     12-1 risk-adj momentum, alignment v3, pillar composite (P1/P2/P3 —
     RED when composite < 6.5 core-eligibility), latest quarterly surprise,
     forward CAGRs, beta, last price                            (automatic)
  4. Technical chart       — 12-month price, 35-day EMA (gold), 200-day MA
     (dashed, the model's trend anchor), ±1σ extension band, annotated with
     the model's trend call / momentum / σ-distance             (automatic)

House style per CLAUDE.md §14: Calibri, navy #1F3A5F, gold #C9A84C, US Letter,
charts at 200 DPI. Footer marks the document internal-use-only.

The quantitative sections are fully automatic from Supabase. The prose is
deliberately NOT automated — supply it per run via --notes (JSON:
{"TICKER": {"biz": "...", "cat": "..."}}); missing entries render as
[TO FILL] placeholders so a draft can go out for prose completion.

Usage:
    python scripts/trade_briefing.py LRCX AMAT LLY --notes notes.json
    python scripts/trade_briefing.py V STX --title "Rotation 2026-08-12" \
        --intro "Custom purpose paragraph..." --out outputs/exports/foo.docx

Example notes file: scripts/briefing_notes_example.json
"""
import sys
import json
import argparse
from datetime import date
from pathlib import Path

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

import psycopg2
from config.settings import settings

NAVY, GOLD, GREY, RED, GREEN = "1F3A5F", "C9A84C", "666666", "B54334", "1E7B4F"


# ── data ──────────────────────────────────────────────────────────────────────
def fetch_model(tickers):
    c = psycopg2.connect(settings.DATABASE_URL); cur = c.cursor()
    out = {}
    for t in tickers:
        cur.execute("""SELECT COALESCE(c.company_name,c.ticker), COALESCE(c.sector,'—'),
            cmd.quadrant, cmd.qgs_tier, cmd.quality_growth_score, cmd.earnings_quality_flag,
            cmd.trend_status, cmd.extension_flag, cmd.mom_12_1_risk_adj, cmd.alignment_score_v3,
            cmd.eps_surprise_q, cmd.rev_surprise_q, cmd.fwd_revenue_3y_cagr, cmd.fwd_eps_3y_cagr,
            cmd.beta, cmd.current_price
            FROM companies c JOIN company_market_data cmd ON cmd.ticker=c.ticker
            AND cmd.data_date=(SELECT MAX(data_date) FROM company_market_data WHERE ticker=c.ticker)
            WHERE c.ticker=%s""", (t,))
        r = cur.fetchone()
        if not r:
            print(f"  ⚠️  {t}: not found in companies — skipped"); continue
        cur.execute("""SELECT s.p1_business_quality,s.p2_management,s.p3_financial_strength,
            s.composite_score_v2 FROM company_scores s JOIN companies c ON c.id=s.company_id
            WHERE c.ticker=%s ORDER BY s.score_date DESC LIMIT 1""", (t,))
        p = cur.fetchone() or (None,) * 4
        f = lambda v: float(v) if v is not None else None
        out[t] = dict(name=r[0], sector=r[1], quad=r[2], tier=r[3], qgs=f(r[4]), eq=r[5],
                      trend=r[6], ext=r[7], mom=f(r[8]), v3=f(r[9]), eps_s=f(r[10]), rev_s=f(r[11]),
                      fwd_rev=f(r[12]), fwd_eps=f(r[13]), beta=f(r[14]), price=f(r[15]),
                      p1=f(p[0]), p2=f(p[1]), p3=f(p[2]), comp=f(p[3]))
    c.close()
    return out


def make_chart(t, m, outdir):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    import pandas as pd
    c = psycopg2.connect(settings.DATABASE_URL); cur = c.cursor()
    cur.execute("""SELECT price_date, adj_close FROM ic_price_history
        WHERE ticker=%s AND adj_close IS NOT NULL ORDER BY price_date""", (t,))
    df = pd.DataFrame(cur.fetchall(), columns=["d", "p"]); c.close()
    if len(df) < 60:
        return None
    df["d"] = pd.to_datetime(df["d"]); df["p"] = df["p"].astype(float)
    df = df.set_index("d")
    df["ema35"] = df["p"].ewm(span=35, adjust=False).mean()
    df["sma200"] = df["p"].rolling(200).mean()
    df["sd200"] = df["p"].rolling(200).std()
    w = df.loc[df.index >= df.index.max() - pd.Timedelta(days=365)]
    ext_z = ((w["p"].iloc[-1] - w["sma200"].iloc[-1]) / w["sd200"].iloc[-1]
             if pd.notna(w["sd200"].iloc[-1]) else float("nan"))
    fig, ax = plt.subplots(figsize=(9.6, 3.6), dpi=200)
    ax.fill_between(w.index, w["sma200"] - w["sd200"], w["sma200"] + w["sd200"],
                    color="#8b95a1", alpha=0.13, linewidth=0, label="200-DMA ±1σ (extension band)")
    ax.plot(w.index, w["p"], color="#1F3A5F", lw=1.6, label="Price")
    ax.plot(w.index, w["ema35"], color="#C9A84C", lw=1.5, label="35-day EMA")
    ax.plot(w.index, w["sma200"], color="#8b95a1", lw=1.2, ls="--", label="200-day MA")
    ax.scatter([w.index[-1]], [w["p"].iloc[-1]], color="#1F3A5F", s=18, zorder=5)
    trend = m["trend"] or "—"
    tcol = "#1E7B4F" if trend == "UPTREND" else ("#B54334" if trend == "DOWNTREND" else "#8b95a1")
    mom = f"{m['mom']:+.2f}" if m["mom"] is not None else "—"
    ax.set_title(f"{t} — {m['name']}", loc="left", fontsize=11, fontweight="bold", color="#1F3A5F")
    ax.text(1.0, 1.06, f"Model trend: {trend}  ·  12-1 risk-adj: {mom}  ·  extension {ext_z:+.1f}σ vs 200-DMA",
            transform=ax.transAxes, ha="right", fontsize=8.5, color=tcol)
    ax.legend(loc="upper left", fontsize=7.5, frameon=False, ncol=4)
    ax.grid(alpha=0.18, lw=0.5)
    ax.spines[["top", "right"]].set_visible(False)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %y"))
    ax.tick_params(labelsize=8)
    ax.yaxis.set_major_formatter(lambda v, _: f"${v:,.0f}")
    fig.tight_layout()
    path = Path(outdir) / f"{t}.png"
    fig.savefig(path, facecolor="white")
    plt.close(fig)
    return path


# ── document ──────────────────────────────────────────────────────────────────
def build_doc(tickers, model, notes, title, intro, out_path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def rgb(h): return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def run(p, text, size=10.5, color="222222", bold=False, italic=False):
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(size)
        r.font.color.rgb = rgb(color); r.bold = bold; r.italic = italic
        return r

    def shade(cell, fill):
        el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(el)

    def bottom_border(p, color=GOLD, sz=12):
        pPr = p._p.get_or_add_pPr(); b = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
        bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), str(sz)); bt.set(qn("w:color"), color)
        b.append(bt); pPr.append(b)

    fmt = lambda v, d=1, suf="": "—" if v is None else f"{v:.{d}f}{suf}"

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    for m_ in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m_, Inches(0.7))

    p = doc.add_paragraph(); run(p, "INTEGRITY COMPOUNDERS", 17, NAVY, bold=True)
    p = doc.add_paragraph(); run(p, title, 13, GOLD, bold=True)
    p = doc.add_paragraph()
    run(p, f"Integrity Wealth Partners · LPL Financial Affiliate  |  For Internal Advisor Use Only  |  "
           f"{date.today().strftime('%B %d, %Y')}  ·  model data as of latest snapshot", 9, GREY)
    bottom_border(p)

    p = doc.add_paragraph(); run(p, "Purpose.  ", bold=True); run(p, intro)
    p = doc.add_paragraph()
    run(p, "Reading the model fields.  ", bold=True)
    run(p, "Stock Quad is the V12 revenue/earnings-momentum quadrant (Q1 best — both accelerating; "
           "Q2 revenue accelerating). QGS is the Quality-Growth Score tier across the universe. "
           "Pillar Composite is the 3-pillar fundamental score (≥8 Tier 1, 6.5–7.9 Tier 2; below 6.5 "
           "is not core-eligible under model rules and is flagged in red). Earnings Quality is the "
           "gross-profit contamination check (EPS_CONFIRMED = earnings growth backed by gross profit). "
           "Charts: 12-month price with 35-day EMA (gold, short-term trend), 200-day MA (dashed, the "
           "model's trend anchor) and the model's ±1σ extension band.")

    chart_dir = Path(out_path).parent / "_briefing_charts"
    chart_dir.mkdir(parents=True, exist_ok=True)
    off_screen = []

    for t in tickers:
        if t not in model:
            continue
        d = model[t]
        n = notes.get(t, {})
        if d["quad"] is None:
            off_screen.append(t)

        p = doc.add_paragraph()
        run(p, f"{t} — {d['name']}", 13, NAVY, bold=True)
        run(p, f"   {d['sector']}", 9.5, GREY)
        bottom_border(p, color="D5DBE3", sz=6)

        p = doc.add_paragraph(); run(p, "What they do.  ", bold=True, color=NAVY)
        run(p, n.get("biz", "[TO FILL — 2-sentence business description]"))
        p = doc.add_paragraph(); run(p, "Catalysts & themes.  ", bold=True, color=NAVY)
        run(p, n.get("cat", "[TO FILL — 2-sentence catalysts / themes]"))

        rows = [
            ("Stock Quad (V12)", d["quad"] or "n/a*",
             "QGS Tier (score)", f"{d['tier']} ({fmt(d['qgs'],5)})" if d["tier"] else "n/a*"),
            ("Earnings Quality", d["eq"] or "n/a*",
             "Model Trend / 12-1 risk-adj", f"{d['trend'] or '—'} / {fmt(d['mom'],2)}"),
            ("Alignment v3", fmt(d["v3"], 0),
             "Pillar Composite (P1/P2/P3)", f"{fmt(d['comp'],2)}  ({fmt(d['p1'])}/{fmt(d['p2'])}/{fmt(d['p3'])})"),
            ("Latest Qtr Surprise (EPS/Rev)",
             f"{fmt(d['eps_s'])}% / {fmt(d['rev_s'])}%" if d["eps_s"] is not None else "n/a*",
             "Fwd 3Y CAGR (Rev/EPS)",
             f"{fmt(d['fwd_rev'])}% / {fmt(d['fwd_eps'])}%"
             + (" †" if (d['fwd_eps'] is not None and abs(d['fwd_eps']) > 300) or (d['fwd_rev'] is not None and abs(d['fwd_rev']) > 300) else "")),
            ("Beta", fmt(d["beta"], 2), "Last Price", f"${d['price']:,.2f}" if d["price"] else "—"),
        ]
        tbl = doc.add_table(rows=len(rows), cols=4)
        tbl.style = "Table Grid"
        for i, (l1, v1, l2, v2) in enumerate(rows):
            for j, val in enumerate((l1, v1, l2, v2)):
                cell = tbl.cell(i, j)
                cell.text = ""
                cp = cell.paragraphs[0]
                is_label = j % 2 == 0
                color = NAVY if is_label else "222222"
                if not is_label:
                    lbl = (l1, v1, l2, v2)[j - 1]
                    if lbl == "Pillar Composite (P1/P2/P3)":
                        color = RED if (d["comp"] is not None and d["comp"] < 6.5) else GREEN
                    if lbl == "Model Trend / 12-1 risk-adj":
                        color = GREEN if d["trend"] == "UPTREND" else (RED if d["trend"] == "DOWNTREND" else GREY)
                run(cp, str(val), 9.5, color, bold=is_label)
                shade(cell, "EFF2F6" if is_label else "FFFFFF")

        img = make_chart(t, d, chart_dir)
        if img:
            p = doc.add_paragraph()
            p.add_run().add_picture(str(img), width=Inches(6.6))
            p = doc.add_paragraph()
            run(p, "12-month price with 35-day EMA (gold), 200-day MA (dashed) and the model's ±1σ "
                   "extension band; header shows the model's trend call, 12-1 risk-adjusted momentum "
                   "and σ-distance from the 200-DMA.", 7.5, GREY, italic=True)
        doc.add_paragraph()

    if off_screen:
        p = doc.add_paragraph()
        run(p, f"* {', '.join(off_screen)}: outside the current Fiscal AI screen — screen-derived "
               "fields (quad, QGS tier, earnings quality, quarterly surprise) not computed this "
               "cycle; pillar scores and momentum are from the model's last full evaluation.",
            8.5, GREY, italic=True)
    p = doc.add_paragraph()
    run(p, "† Forward CAGR beyond ±300% is a base-effect artifact of near-zero trailing figures — "
           "treat as directional only. ", 8.5, GREY, italic=True)
    p = doc.add_paragraph()
    run(p, "Source: Integrity Compounders Alpha System v12.1. Company facts from public reporting. "
           "For internal use by Integrity Wealth Partners advisors only — not for client "
           "distribution, not investment advice. Model output reflects systematic signals and does "
           "not account for client suitability.", 8.5, GREY, italic=True)

    doc.save(out_path)
    print(f"  ✅ {out_path}")


def main():
    ap = argparse.ArgumentParser(description="Generate an advisor trade briefing (house template)")
    ap.add_argument("tickers", nargs="+", help="tickers to include, in order")
    ap.add_argument("--notes", help="JSON file: {TICKER: {biz: ..., cat: ...}}")
    ap.add_argument("--title", default="Model Additions Briefing")
    ap.add_argument("--intro", default=("Client portfolios tracking the Integrity Compounders model "
                    "are missing several current holdings; this briefing covers the targeted adds. "
                    "The quantitative model output is shown for each name so the full picture — "
                    "including where systematic scores run below core thresholds — is on the table."))
    ap.add_argument("--out", default=None)
    a = ap.parse_args()
    tickers = [t.upper() for t in a.tickers]
    notes = json.load(open(a.notes)) if a.notes else {}
    out = a.out or str(ROOT / "outputs" / "exports" / f"trade_briefing_{date.today()}.docx")
    model = fetch_model(tickers)
    build_doc(tickers, model, notes, a.title, a.intro, out)


if __name__ == "__main__":
    main()
